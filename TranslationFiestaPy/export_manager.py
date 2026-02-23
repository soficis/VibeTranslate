#!/usr/bin/env python3
"""
TranslationFiesta Export Manager
Provides professional document export capabilities with proper formatting,
metadata inclusion, and quality metrics.
"""

import os
from dataclasses import asdict, dataclass
from datetime import datetime
from typing import List, Optional

# Export dependencies
try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from jinja2 import Environment, FileSystemLoader
    JINJA2_AVAILABLE = True
except ImportError:
    JINJA2_AVAILABLE = False

from app_logger import create_logger
from bleu_scorer import BLEUScorer
from exceptions import TranslationFiestaError


@dataclass
class ExportMetadata:
    """Metadata for exported documents"""
    title: str
    author: str = "TranslationFiesta"
    subject: str = "Translation Results"
    keywords: List[str] = None
    created_date: str = None
    source_language: str = ""
    target_language: str = ""
    translation_quality_score: float = 0.0
    processing_time_seconds: float = 0.0
    api_used: str = ""

    def __post_init__(self):
        if self.keywords is None:
            self.keywords = ["translation", "localization"]
        if self.created_date is None:
            self.created_date = datetime.now().isoformat()


@dataclass
class ExportConfig:
    """Configuration for document export"""
    format: str = "pdf"  # pdf, docx, html
    template_path: Optional[str] = None
    include_metadata: bool = True
    include_quality_metrics: bool = True
    include_timestamps: bool = True
    page_size: str = "A4"  # A4, letter
    font_family: str = "Helvetica"
    font_size: int = 12
    include_table_of_contents: bool = False
    compress_output: bool = False
    custom_css: Optional[str] = None


@dataclass
class TranslationResult:
    """Represents a translation result for export"""
    original_text: str
    translated_text: str
    source_language: str
    target_language: str
    quality_score: float = 0.0
    confidence_level: str = ""
    processing_time: float = 0.0
    api_used: str = ""
    timestamp: str = ""

    def __post_init__(self):
        if not self.timestamp:
            self.timestamp = datetime.now().isoformat()


class ExportManager:
    """Main export manager for TranslationFiesta"""

    SUPPORTED_FORMATS = ["pdf", "docx", "html"]

    def __init__(self, config: Optional[ExportConfig] = None):
        self.config = config or ExportConfig()
        self.logger = create_logger("export_manager")
        self.bleu_scorer = BLEUScorer()

        # Check dependencies
        if self.config.format == "pdf" and not REPORTLAB_AVAILABLE:
            raise TranslationFiestaError("PDF export requires 'reportlab' package. Install with: pip install reportlab")
        if self.config.format == "docx" and not DOCX_AVAILABLE:
            raise TranslationFiestaError("DOCX export requires 'python-docx' package. Install with: pip install python-docx")
        if self.config.template_path and not JINJA2_AVAILABLE:
            raise TranslationFiestaError("Template support requires 'jinja2' package. Install with: pip install jinja2")

    def export_translations(
        self,
        translations: List[TranslationResult],
        output_path: str,
        metadata: Optional[ExportMetadata] = None
    ) -> str:
        """
        Export translations to the specified format

        Args:
            translations: List of translation results
            output_path: Path to save the exported file
            metadata: Optional metadata for the document

        Returns:
            Path to the exported file
        """
        if not translations:
            raise TranslationFiestaError("No translations provided for export")

        # Create default metadata if not provided
        if metadata is None:
            metadata = ExportMetadata(
                title=f"Translation Results - {len(translations)} items",
                source_language=translations[0].source_language if translations else "",
                target_language=translations[0].target_language if translations else "",
            )

        # Calculate overall quality metrics
        if self.config.include_quality_metrics:
            self._calculate_quality_metrics(translations, metadata)

        # Export based on format
        if self.config.format == "pdf":
            return self._export_pdf(translations, output_path, metadata)
        elif self.config.format == "docx":
            return self._export_docx(translations, output_path, metadata)
        elif self.config.format == "html":
            return self._export_html(translations, output_path, metadata)
        else:
            raise TranslationFiestaError(f"Unsupported export format: {self.config.format}")

    def _calculate_quality_metrics(self, translations: List[TranslationResult], metadata: ExportMetadata):
        """Calculate overall quality metrics for the translation set"""
        total_score = 0.0
        total_time = 0.0

        for translation in translations:
            if translation.quality_score == 0.0 and translation.original_text and translation.translated_text:
                # Calculate BLEU score if not already calculated
                bleu_score = self.bleu_scorer.calculate_bleu(translation.original_text, translation.translated_text)
                translation.quality_score = bleu_score
                confidence = self.bleu_scorer.get_confidence_level(bleu_score)
                translation.confidence_level = confidence["Level"]

            total_score += translation.quality_score
            total_time += translation.processing_time

        # Update metadata with averages
        if translations:
            metadata.translation_quality_score = total_score / len(translations)
            metadata.processing_time_seconds = total_time / len(translations)

    def _export_pdf(
        self,
        translations: List[TranslationResult],
        output_path: str,
        metadata: ExportMetadata
    ) -> str:
        """Export translations to PDF format"""
        if not REPORTLAB_AVAILABLE:
            raise TranslationFiestaError("PDF export not available - missing reportlab dependency")

        # Create PDF document
        doc = SimpleDocTemplate(
            output_path,
            pagesize=A4 if self.config.page_size == "A4" else letter,
            title=metadata.title,
            author=metadata.author,
            subject=metadata.subject,
            keywords=", ".join(metadata.keywords)
        )

        # Get styles
        styles = getSampleStyleSheet()
        title_style = styles["Title"]
        heading_style = styles["Heading2"]
        normal_style = styles["Normal"]

        # Create content
        content = []

        # Title
        content.append(Paragraph(metadata.title, title_style))
        content.append(Spacer(1, 0.5 * inch))

        # Metadata table
        if self.config.include_metadata:
            metadata_data = self._create_metadata_table_data(metadata)
            metadata_table = Table(metadata_data)
            metadata_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 14),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            content.append(metadata_table)
            content.append(Spacer(1, 0.5 * inch))

        # Translations
        content.append(Paragraph("Translation Results", heading_style))
        content.append(Spacer(1, 0.25 * inch))

        for i, translation in enumerate(translations, 1):
            # Translation header
            content.append(Paragraph(f"Translation {i}", styles["Heading3"]))

            # Original text
            content.append(Paragraph("<b>Original Text:</b>", normal_style))
            content.append(Paragraph(translation.original_text, normal_style))

            # Translated text
            content.append(Paragraph("<b>Translated Text:</b>", normal_style))
            content.append(Paragraph(translation.translated_text, normal_style))

            # Quality metrics
            if self.config.include_quality_metrics:
                quality_text = f"Quality Score: {translation.quality_score:.3f} ({translation.confidence_level})"
                if translation.processing_time > 0:
                    quality_text += f" | Processing Time: {translation.processing_time:.2f}s"
                content.append(Paragraph(f"<i>{quality_text}</i>", normal_style))

            content.append(Spacer(1, 0.25 * inch))

        # Build PDF
        doc.build(content)
        return output_path

    def _export_docx(
        self,
        translations: List[TranslationResult],
        output_path: str,
        metadata: ExportMetadata
    ) -> str:
        """Export translations to DOCX format"""
        if not DOCX_AVAILABLE:
            raise TranslationFiestaError("DOCX export not available - missing python-docx dependency")

        # Create document
        doc = Document()

        # Set document properties
        doc.core_properties.title = metadata.title
        doc.core_properties.author = metadata.author
        doc.core_properties.subject = metadata.subject
        doc.core_properties.keywords = ", ".join(metadata.keywords)
        doc.core_properties.created = datetime.fromisoformat(metadata.created_date)

        # Title
        title = doc.add_heading(metadata.title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata table
        if self.config.include_metadata:
            self._add_metadata_table_docx(doc, metadata)

        # Translations section
        doc.add_heading("Translation Results", 1)

        for i, translation in enumerate(translations, 1):
            # Translation header
            doc.add_heading(f"Translation {i}", 2)

            # Original text
            doc.add_paragraph("Original Text:", style='Intense Quote')
            p = doc.add_paragraph(translation.original_text)
            p.style = 'Body Text'

            # Translated text
            doc.add_paragraph("Translated Text:", style='Intense Quote')
            p = doc.add_paragraph(translation.translated_text)
            p.style = 'Body Text'

            # Quality metrics
            if self.config.include_quality_metrics:
                quality_text = f"Quality Score: {translation.quality_score:.3f} ({translation.confidence_level})"
                if translation.processing_time > 0:
                    quality_text += f" | Processing Time: {translation.processing_time:.2f}s"
                p = doc.add_paragraph(quality_text, style='Caption')
                p.italic = True

            # Add spacing
            doc.add_paragraph("")

        # Save document
        doc.save(output_path)
        return output_path

    def _export_html(
        self,
        translations: List[TranslationResult],
        output_path: str,
        metadata: ExportMetadata
    ) -> str:
        """Export translations to HTML format"""
        # Use template if available
        if self.config.template_path and JINJA2_AVAILABLE:
            return self._export_html_with_template(translations, output_path, metadata)
        else:
            return self._export_html_basic(translations, output_path, metadata)

    def _export_html_basic(
        self,
        translations: List[TranslationResult],
        output_path: str,
        metadata: ExportMetadata
    ) -> str:
        """Export to basic HTML format"""
        html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{metadata.title}</title>
    <style>
        body {{
            font-family: {self.config.font_family}, sans-serif;
            font-size: {self.config.font_size}px;
            line-height: 1.6;
            margin: 40px;
            background-color: #f5f5f5;
        }}
        .container {{
            max-width: 800px;
            margin: 0 auto;
            background: white;
            padding: 30px;
            border-radius: 8px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        }}
        h1 {{
            color: #333;
            text-align: center;
            border-bottom: 2px solid #007acc;
            padding-bottom: 10px;
        }}
        h2 {{
            color: #555;
            margin-top: 30px;
        }}
        .translation {{
            margin-bottom: 30px;
            padding: 20px;
            background: #f9f9f9;
            border-left: 4px solid #007acc;
        }}
        .original {{
            margin-bottom: 15px;
        }}
        .translated {{
            margin-bottom: 15px;
        }}
        .quality {{
            font-style: italic;
            color: #666;
            font-size: 0.9em;
        }}
        .metadata {{
            background: #e8f4fd;
            padding: 15px;
            border-radius: 5px;
            margin-bottom: 20px;
        }}
        table {{
            width: 100%;
            border-collapse: collapse;
        }}
        th, td {{
            padding: 8px 12px;
            text-align: left;
            border-bottom: 1px solid #ddd;
        }}
        th {{
            background-color: #007acc;
            color: white;
        }}
    </style>
    {self.config.custom_css or ""}
</head>
<body>
    <div class="container">
        <h1>{metadata.title}</h1>
"""

        # Metadata
        if self.config.include_metadata:
            html_content += f"""
        <div class="metadata">
            <h2>Document Information</h2>
            <table>
                <tr><th>Author</th><td>{metadata.author}</td></tr>
                <tr><th>Created</th><td>{metadata.created_date}</td></tr>
                <tr><th>Source Language</th><td>{metadata.source_language}</td></tr>
                <tr><th>Target Language</th><td>{metadata.target_language}</td></tr>
                <tr><th>Quality Score</th><td>{metadata.translation_quality_score:.3f}</td></tr>
                <tr><th>API Used</th><td>{metadata.api_used}</td></tr>
            </table>
        </div>
"""

        # Translations
        html_content += "<h2>Translation Results</h2>"

        for i, translation in enumerate(translations, 1):
            quality_info = ""
            if self.config.include_quality_metrics:
                quality_info = f'<div class="quality">Quality Score: {translation.quality_score:.3f} ({translation.confidence_level})'
                if translation.processing_time > 0:
                    quality_info += f' | Processing Time: {translation.processing_time:.2f}s'
                quality_info += '</div>'

            html_content += f"""
        <div class="translation">
            <h3>Translation {i}</h3>
            <div class="original">
                <strong>Original Text:</strong><br>
                {translation.original_text.replace(chr(10), '<br>')}
            </div>
            <div class="translated">
                <strong>Translated Text:</strong><br>
                {translation.translated_text.replace(chr(10), '<br>')}
            </div>
            {quality_info}
        </div>
"""

        html_content += """
    </div>
</body>
</html>"""

        # Write to file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)

        return output_path

    def _export_html_with_template(
        self,
        translations: List[TranslationResult],
        output_path: str,
        metadata: ExportMetadata
    ) -> str:
        """Export HTML using Jinja2 template"""
        if not JINJA2_AVAILABLE:
            raise TranslationFiestaError("Template support not available - missing jinja2 dependency")

        template_dir = os.path.dirname(self.config.template_path)
        template_name = os.path.basename(self.config.template_path)

        env = Environment(loader=FileSystemLoader(template_dir))
        template = env.get_template(template_name)

        # Prepare template data
        template_data = {
            'metadata': asdict(metadata),
            'translations': [asdict(t) for t in translations],
            'config': asdict(self.config),
            'timestamp': datetime.now().isoformat()
        }

        # Render template
        html_content = template.render(**template_data)

        # Write to file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)

        return output_path

    def _create_metadata_table_data(self, metadata: ExportMetadata) -> List[List[str]]:
        """Create metadata table data for PDF export"""
        return [
            ["Property", "Value"],
            ["Title", metadata.title],
            ["Author", metadata.author],
            ["Created", metadata.created_date],
            ["Source Language", metadata.source_language],
            ["Target Language", metadata.target_language],
            ["Quality Score", ".3f"],
            ["API Used", metadata.api_used],
            ["Processing Time", ".2f"]
        ]

    def _add_metadata_table_docx(self, doc: Document, metadata: ExportMetadata):
        """Add metadata table to DOCX document"""
        table = doc.add_table(rows=8, cols=2)
        table.style = 'Table Grid'

        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Property'
        hdr_cells[1].text = 'Value'

        # Data rows
        data = [
            ("Title", metadata.title),
            ("Author", metadata.author),
            ("Created", metadata.created_date),
            ("Source Language", metadata.source_language),
            ("Target Language", metadata.target_language),
            ("Quality Score", ".3f"),
            ("API Used", metadata.api_used),
            ("Processing Time", ".2f")
        ]

        for i, (prop, value) in enumerate(data, 1):
            row_cells = table.rows[i].cells
            row_cells[0].text = prop
            row_cells[1].text = str(value)

    def create_template(self, template_path: str, template_type: str = "html"):
        """Create a default template file"""
        if template_type == "html":
            template_content = """<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{{ metadata.title }}</title>
    <style>
        body { font-family: Arial, sans-serif; margin: 40px; }
        .header { text-align: center; border-bottom: 2px solid #333; padding-bottom: 20px; }
        .translation { margin: 20px 0; padding: 15px; border: 1px solid #ddd; }
        .original { background: #f0f0f0; padding: 10px; }
        .translated { background: #e8f4fd; padding: 10px; }
        .quality { font-style: italic; color: #666; }
    </style>
</head>
<body>
    <div class="header">
        <h1>{{ metadata.title }}</h1>
        <p>Generated on {{ timestamp }}</p>
    </div>

    {% if config.include_metadata %}
    <div class="metadata">
        <h2>Document Information</h2>
        <ul>
            <li><strong>Author:</strong> {{ metadata.author }}</li>
            <li><strong>Source Language:</strong> {{ metadata.source_language }}</li>
            <li><strong>Target Language:</strong> {{ metadata.target_language }}</li>
            <li><strong>Quality Score:</strong> {{ metadata.translation_quality_score | round(3) }}</li>
        </ul>
    </div>
    {% endif %}

    <h2>Translations</h2>
    {% for translation in translations %}
    <div class="translation">
        <div class="original">
            <strong>Original:</strong><br>
            {{ translation.original_text | replace('\n', '<br>') | safe }}
        </div>
        <div class="translated">
            <strong>Translated:</strong><br>
            {{ translation.translated_text | replace('\n', '<br>') | safe }}
        </div>
        {% if config.include_quality_metrics %}
        <div class="quality">
            Quality Score: {{ translation.quality_score | round(3) }} ({{ translation.confidence_level }})
            {% if translation.processing_time > 0 %}
            | Processing Time: {{ translation.processing_time | round(2) }}s
            {% endif %}
        </div>
        {% endif %}
    </div>
    {% endfor %}
</body>
</html>"""
        else:
            raise TranslationFiestaError(f"Unsupported template type: {template_type}")

        # Write template
        os.makedirs(os.path.dirname(template_path), exist_ok=True)
        with open(template_path, 'w', encoding='utf-8') as f:
            f.write(template_content)

        return template_path


# Convenience functions
def export_to_pdf(
    translations: List[TranslationResult],
    output_path: str,
    metadata: Optional[ExportMetadata] = None,
    config: Optional[ExportConfig] = None
) -> str:
    """Convenience function to export to PDF"""
    config = config or ExportConfig(format="pdf")
    config.format = "pdf"
    manager = ExportManager(config)
    return manager.export_translations(translations, output_path, metadata)


def export_to_docx(
    translations: List[TranslationResult],
    output_path: str,
    metadata: Optional[ExportMetadata] = None,
    config: Optional[ExportConfig] = None
) -> str:
    """Convenience function to export to DOCX"""
    config = config or ExportConfig(format="docx")
    config.format = "docx"
    manager = ExportManager(config)
    return manager.export_translations(translations, output_path, metadata)


def export_to_html(
    translations: List[TranslationResult],
    output_path: str,
    metadata: Optional[ExportMetadata] = None,
    config: Optional[ExportConfig] = None
) -> str:
    """Convenience function to export to HTML"""
    config = config or ExportConfig(format="html")
    config.format = "html"
    manager = ExportManager(config)
    return manager.export_translations(translations, output_path, metadata)


if __name__ == "__main__":
    # Example usage
    translations = [
        TranslationResult(
            original_text="Hello, how are you?",
            translated_text="こんにちは、お元気ですか？",
            source_language="en",
            target_language="ja",
            quality_score=0.85,
            confidence_level="High",
            processing_time=1.2,
            api_used="Google Translate"
        ),
        TranslationResult(
            original_text="Thank you for your help.",
            translated_text="お手伝いいただきありがとうございます。",
            source_language="en",
            target_language="ja",
            quality_score=0.92,
            confidence_level="Very High",
            processing_time=0.8,
            api_used="Google Translate"
        )
    ]

    metadata = ExportMetadata(
        title="Sample Translation Export",
        source_language="en",
        target_language="ja",
        api_used="Google Translate"
    )

    # Export to different formats
    try:
        pdf_path = export_to_pdf(translations, "sample_translations.pdf", metadata)
        print(f"PDF exported to: {pdf_path}")
    except Exception as e:
        print(f"PDF export failed: {e}")

    try:
        docx_path = export_to_docx(translations, "sample_translations.docx", metadata)
        print(f"DOCX exported to: {docx_path}")
    except Exception as e:
        print(f"DOCX export failed: {e}")

    try:
        html_path = export_to_html(translations, "sample_translations.html", metadata)
        print(f"HTML exported to: {html_path}")
    except Exception as e:
        print(f"HTML export failed: {e}")
